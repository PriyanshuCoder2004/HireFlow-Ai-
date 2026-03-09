from fastapi import FastAPI, APIRouter, HTTPException, Depends, UploadFile, File, status, Form, BackgroundTasks
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
import os
import logging
import io
import re
import tempfile
import asyncio
from pathlib import Path
from pydantic import BaseModel, Field, EmailStr
from typing import List, Optional
import uuid
from datetime import datetime, timezone, timedelta
import jwt
import bcrypt
from emergentintegrations.llm.chat import LlmChat, UserMessage
import PyPDF2
from docx import Document
import pytesseract
from pdf2image import convert_from_bytes
from PIL import Image
import resend
from apscheduler.schedulers.asyncio import AsyncIOScheduler
from apscheduler.triggers.interval import IntervalTrigger

ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

# MongoDB connection
mongo_url = os.environ['MONGO_URL']
client = AsyncIOMotorClient(mongo_url)
db = client[os.environ['DB_NAME']]

# JWT Configuration
JWT_SECRET = os.environ.get('JWT_SECRET', 'hireflow_ai_secret_key_2025')
JWT_ALGORITHM = "HS256"
JWT_EXPIRATION_HOURS = 24

# LLM Key
EMERGENT_LLM_KEY = os.environ.get('EMERGENT_LLM_KEY')

# Email Configuration (Resend)
RESEND_API_KEY = os.environ.get('RESEND_API_KEY')
SENDER_EMAIL = os.environ.get('SENDER_EMAIL', 'onboarding@resend.dev')
if RESEND_API_KEY:
    resend.api_key = RESEND_API_KEY

app = FastAPI(title="HireFlow AI API")
api_router = APIRouter(prefix="/api")
security = HTTPBearer()

# Background Scheduler for email reminders
scheduler = AsyncIOScheduler()

# Debug mode for development
DEBUG_MODE = os.environ.get('DEBUG_MODE', 'true').lower() == 'true'

# Scheduler status tracking (for debug endpoint)
scheduler_status = {
    "last_run_time": None,
    "last_run_duration_ms": None,
    "total_events_checked": 0,
    "eligible_24hr_count": 0,
    "eligible_1hr_count": 0,
    "reminders_sent_24hr": 0,
    "reminders_sent_1hr": 0,
    "last_errors": [],
    "execution_logs": []
}

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# ===================== MODELS =====================

class UserCreate(BaseModel):
    email: EmailStr
    password: str
    name: str

class UserLogin(BaseModel):
    email: EmailStr
    password: str

class UserResponse(BaseModel):
    id: str
    email: str
    name: str
    created_at: str

class TokenResponse(BaseModel):
    token: str
    user: UserResponse

class ResumeCreate(BaseModel):
    content: str
    title: str

class ResumeResponse(BaseModel):
    id: str
    user_id: str
    title: str
    content: str
    file_name: Optional[str] = None
    file_type: Optional[str] = None
    extraction_method: Optional[str] = None  # "parser" or "ocr"
    extraction_status: Optional[str] = None  # "success" or "partial" or "failed"
    ocr_used: bool = False
    analysis: Optional[dict] = None
    score: Optional[int] = None
    created_at: str
    updated_at: str

class JobApplicationCreate(BaseModel):
    company: str
    position: str
    job_url: Optional[str] = None
    job_description: Optional[str] = None
    status: str = "applied"
    notes: Optional[str] = None
    applied_date: Optional[str] = None

class JobApplicationUpdate(BaseModel):
    company: Optional[str] = None
    position: Optional[str] = None
    job_url: Optional[str] = None
    job_description: Optional[str] = None
    status: Optional[str] = None
    notes: Optional[str] = None

class JobApplicationResponse(BaseModel):
    id: str
    user_id: str
    company: str
    position: str
    job_url: Optional[str] = None
    job_description: Optional[str] = None
    status: str
    notes: Optional[str] = None
    applied_date: str
    created_at: str
    updated_at: str

class CoverLetterRequest(BaseModel):
    job_description: str
    company_name: str
    position: str
    resume_id: Optional[str] = None

class CoverLetterResponse(BaseModel):
    id: str
    user_id: str
    company_name: str
    position: str
    content: str
    created_at: str

# Phase 1 Cover Letter Generator Models
class CoverLetterGenerateRequest(BaseModel):
    resume_id: str
    job_application_id: str
    customization_notes: Optional[str] = None

class CoverLetterUpdateRequest(BaseModel):
    content: Optional[str] = None
    title: Optional[str] = None

class CoverLetterFullResponse(BaseModel):
    id: str
    user_id: str
    resume_id: str
    job_application_id: str
    title: str
    content: str
    word_count: int
    company_name: str
    position: str
    created_at: str
    updated_at: str

class CalendarEventCreate(BaseModel):
    title: str
    description: Optional[str] = None
    event_type: str = "interview"  # interview, phone_screen, video_call, follow_up, other
    interview_type: Optional[str] = None  # hr, technical, managerial, final, other
    start_date: str
    end_date: Optional[str] = None
    job_application_id: Optional[str] = None
    location: Optional[str] = None
    meeting_link: Optional[str] = None
    notes: Optional[str] = None
    reminders_enabled: bool = True
    reminder_24hr_sent: bool = False
    reminder_1hr_sent: bool = False

class CalendarEventUpdate(BaseModel):
    title: Optional[str] = None
    description: Optional[str] = None
    event_type: Optional[str] = None
    interview_type: Optional[str] = None
    start_date: Optional[str] = None
    end_date: Optional[str] = None
    location: Optional[str] = None
    meeting_link: Optional[str] = None
    notes: Optional[str] = None
    reminders_enabled: Optional[bool] = None

class CalendarEventResponse(BaseModel):
    id: str
    user_id: str
    title: str
    description: Optional[str] = None
    event_type: str
    interview_type: Optional[str] = None
    start_date: str
    end_date: Optional[str] = None
    job_application_id: Optional[str] = None
    location: Optional[str] = None
    meeting_link: Optional[str] = None
    notes: Optional[str] = None
    reminders_enabled: bool = True
    reminder_24hr_sent: bool = False
    reminder_1hr_sent: bool = False
    created_at: str

# ===================== NOTIFICATION MODELS =====================

class NotificationLog(BaseModel):
    id: str
    user_id: str
    event_id: str
    job_application_id: Optional[str] = None
    reminder_type: str  # "24hr" or "1hr"
    delivery_status: str  # "sent", "failed", "pending"
    error_message: Optional[str] = None
    sent_timestamp: Optional[str] = None
    created_at: str

class NotificationPreferences(BaseModel):
    email_reminders_enabled: bool = True
    reminder_24hr: bool = True
    reminder_1hr: bool = True

class MatchRequest(BaseModel):
    resume_id: str
    job_description: str
    job_title: Optional[str] = None
    company_name: Optional[str] = None

class MatchAnalysis(BaseModel):
    match_score: int
    skill_match: dict
    experience_match: dict
    missing_skills: List[str]
    weak_areas: List[str]
    strengths: List[str]
    suggestions: List[str]
    keyword_analysis: dict
    summary: str

class JobMatchResponse(BaseModel):
    id: str
    user_id: str
    resume_id: str
    job_title: Optional[str] = None
    company_name: Optional[str] = None
    job_description: str
    analysis: MatchAnalysis
    created_at: str

class MatchResponse(BaseModel):
    match_score: int
    strengths: List[str]
    gaps: List[str]
    suggestions: List[str]

# ===================== INTERVIEW PREP MODELS =====================

class InterviewQuestion(BaseModel):
    question: str
    category: str  # "hr_behavioral", "technical", "scenario"
    difficulty: str  # "easy", "medium", "hard"
    guidance: List[str]  # STAR method hints or bullet points
    sample_points: List[str]  # Key points to cover

class WeakArea(BaseModel):
    topic: str
    reason: str
    preparation_tips: List[str]
    resources: Optional[List[str]] = None

class InterviewPrepAnalysis(BaseModel):
    hr_behavioral_questions: List[InterviewQuestion]
    technical_questions: List[InterviewQuestion]
    scenario_questions: List[InterviewQuestion]
    weak_areas: List[WeakArea]
    general_tips: List[str]
    company_research_points: List[str]
    questions_to_ask: List[str]

class InterviewPrepRequest(BaseModel):
    application_id: str
    resume_id: str
    include_match_analysis: bool = True

class InterviewPrepResponse(BaseModel):
    id: str
    user_id: str
    application_id: str
    resume_id: str
    job_title: str
    company_name: str
    analysis: InterviewPrepAnalysis
    match_score: Optional[int] = None
    ai_generated: Optional[bool] = None
    created_at: str
    updated_at: str

# ===================== HELPERS =====================

def hash_password(password: str) -> str:
    return bcrypt.hashpw(password.encode('utf-8'), bcrypt.gensalt()).decode('utf-8')

def verify_password(password: str, hashed: str) -> bool:
    return bcrypt.checkpw(password.encode('utf-8'), hashed.encode('utf-8'))

def create_token(user_id: str) -> str:
    payload = {
        "user_id": user_id,
        "exp": datetime.now(timezone.utc) + timedelta(hours=JWT_EXPIRATION_HOURS)
    }
    return jwt.encode(payload, JWT_SECRET, algorithm=JWT_ALGORITHM)

async def get_current_user(credentials: HTTPAuthorizationCredentials = Depends(security)):
    try:
        payload = jwt.decode(credentials.credentials, JWT_SECRET, algorithms=[JWT_ALGORITHM])
        user_id = payload.get("user_id")
        if not user_id:
            raise HTTPException(status_code=401, detail="Invalid token")
        user = await db.users.find_one({"id": user_id}, {"_id": 0, "password": 0})
        if not user:
            raise HTTPException(status_code=401, detail="User not found")
        return user
    except jwt.ExpiredSignatureError:
        raise HTTPException(status_code=401, detail="Token expired")
    except jwt.InvalidTokenError:
        raise HTTPException(status_code=401, detail="Invalid token")

async def get_llm_response(system_message: str, user_message: str) -> str:
    try:
        chat = LlmChat(
            api_key=EMERGENT_LLM_KEY,
            session_id=str(uuid.uuid4()),
            system_message=system_message
        ).with_model("openai", "gpt-5.2")
        
        message = UserMessage(text=user_message)
        response = await chat.send_message(message)
        return response
    except Exception as e:
        logger.error(f"LLM error: {e}")
        raise HTTPException(status_code=500, detail="AI service temporarily unavailable")

async def get_llm_response_safe(system_message: str, user_message: str) -> tuple[str, bool]:
    """
    Safe version of get_llm_response that returns (response, success) tuple.
    Never raises exceptions - returns empty string and False on failure.
    """
    try:
        chat = LlmChat(
            api_key=EMERGENT_LLM_KEY,
            session_id=str(uuid.uuid4()),
            system_message=system_message
        ).with_model("openai", "gpt-5.2")
        
        message = UserMessage(text=user_message)
        response = await chat.send_message(message)
        if response and len(response.strip()) > 50:
            return response, True
        return "", False
    except Exception as e:
        logger.error(f"LLM error (safe mode): {e}")
        return "", False

# ===================== FILE TEXT EXTRACTION WITH OCR =====================

# Minimum character threshold for valid text extraction
MIN_TEXT_LENGTH = 100
MIN_WORD_COUNT = 20

class ExtractionResult:
    """Result of text extraction with metadata"""
    def __init__(self, text: str, method: str, status: str, ocr_used: bool = False):
        self.text = text
        self.method = method  # "parser" or "ocr"
        self.status = status  # "success", "partial", "failed"
        self.ocr_used = ocr_used

def is_text_readable(text: str) -> bool:
    """Check if extracted text is readable and meaningful"""
    if not text or len(text.strip()) < MIN_TEXT_LENGTH:
        return False
    
    # Count words (at least MIN_WORD_COUNT readable words)
    words = re.findall(r'\b[a-zA-Z]{2,}\b', text)
    if len(words) < MIN_WORD_COUNT:
        return False
    
    # Check for high ratio of special characters (might indicate garbled text)
    alpha_ratio = len(re.findall(r'[a-zA-Z]', text)) / len(text) if text else 0
    if alpha_ratio < 0.3:  # Less than 30% alphabetic characters
        return False
    
    return True

def extract_text_from_pdf_parser(file_content: bytes) -> str:
    """Extract text content from PDF using PyPDF2 parser"""
    try:
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
        text_content = []
        for page in pdf_reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_content.append(page_text)
        return "\n".join(text_content).strip()
    except Exception as e:
        logger.error(f"PDF parser extraction error: {e}")
        return ""

def extract_text_from_pdf_ocr(file_content: bytes) -> str:
    """Extract text from PDF using OCR (for scanned/image-based PDFs)"""
    try:
        # Convert PDF pages to images
        images = convert_from_bytes(file_content, dpi=300)
        
        text_content = []
        for i, image in enumerate(images):
            # Use Tesseract OCR on each page image
            page_text = pytesseract.image_to_string(image, lang='eng')
            if page_text and page_text.strip():
                text_content.append(page_text.strip())
            logger.info(f"OCR processed page {i+1}/{len(images)}")
        
        return "\n\n".join(text_content).strip()
    except Exception as e:
        logger.error(f"PDF OCR extraction error: {e}")
        return ""

def extract_text_from_image_ocr(file_content: bytes) -> str:
    """Extract text from image file using OCR"""
    try:
        image = Image.open(io.BytesIO(file_content))
        text = pytesseract.image_to_string(image, lang='eng')
        return text.strip()
    except Exception as e:
        logger.error(f"Image OCR extraction error: {e}")
        return ""

def extract_text_from_pdf(file_content: bytes) -> ExtractionResult:
    """Extract text from PDF with OCR fallback"""
    # Step 1: Try standard PDF text extraction
    logger.info("Attempting PDF text extraction with parser...")
    parser_text = extract_text_from_pdf_parser(file_content)
    
    if is_text_readable(parser_text):
        logger.info(f"Parser extraction successful: {len(parser_text)} chars")
        return ExtractionResult(
            text=parser_text,
            method="parser",
            status="success",
            ocr_used=False
        )
    
    # Step 2: Parser failed or text unreadable, try OCR
    logger.info("Parser extraction insufficient, attempting OCR fallback...")
    ocr_text = extract_text_from_pdf_ocr(file_content)
    
    if is_text_readable(ocr_text):
        logger.info(f"OCR extraction successful: {len(ocr_text)} chars")
        return ExtractionResult(
            text=ocr_text,
            method="ocr",
            status="success",
            ocr_used=True
        )
    
    # Step 3: Both methods produced some text but below threshold
    # Return the better result with partial status
    if len(ocr_text) > len(parser_text):
        if ocr_text:
            return ExtractionResult(
                text=ocr_text,
                method="ocr",
                status="partial",
                ocr_used=True
            )
    elif parser_text:
        return ExtractionResult(
            text=parser_text,
            method="parser",
            status="partial",
            ocr_used=False
        )
    
    # Step 4: Complete failure
    return ExtractionResult(
        text="",
        method="failed",
        status="failed",
        ocr_used=True
    )

def extract_text_from_docx(file_content: bytes) -> ExtractionResult:
    """Extract text content from DOCX file"""
    try:
        doc = Document(io.BytesIO(file_content))
        text_content = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    text_content.append(" | ".join(row_text))
        
        text = "\n".join(text_content).strip()
        
        if is_text_readable(text):
            return ExtractionResult(
                text=text,
                method="parser",
                status="success",
                ocr_used=False
            )
        elif text:
            return ExtractionResult(
                text=text,
                method="parser",
                status="partial",
                ocr_used=False
            )
        else:
            return ExtractionResult(
                text="",
                method="failed",
                status="failed",
                ocr_used=False
            )
    except Exception as e:
        logger.error(f"DOCX extraction error: {e}")
        return ExtractionResult(
            text="",
            method="failed",
            status="failed",
            ocr_used=False
        )

# ===================== EMAIL NOTIFICATION SYSTEM =====================

def generate_interview_reminder_email(
    candidate_name: str,
    company_name: str,
    job_role: str,
    interview_date: str,
    interview_time: str,
    interview_type: str,
    location: str,
    meeting_link: str,
    reminder_type: str
) -> tuple:
    """Generate professional interview reminder email"""
    
    time_message = "24 hours" if reminder_type == "24hr" else "1 hour"
    interview_type_display = interview_type.replace("_", " ").title() if interview_type else "Interview"
    
    subject = f"⏰ Reminder: {interview_type_display} at {company_name} in {time_message}"
    
    location_section = ""
    if meeting_link:
        location_section = f"""
        <tr>
            <td style="padding: 8px 0; color: #64748b;">Meeting Link:</td>
            <td style="padding: 8px 0;"><a href="{meeting_link}" style="color: #6366f1;">{meeting_link}</a></td>
        </tr>
        """
    elif location:
        location_section = f"""
        <tr>
            <td style="padding: 8px 0; color: #64748b;">Location:</td>
            <td style="padding: 8px 0;">{location}</td>
        </tr>
        """
    
    html_body = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="utf-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
    </head>
    <body style="margin: 0; padding: 0; font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif; background-color: #f8fafc;">
        <div style="max-width: 600px; margin: 0 auto; padding: 40px 20px;">
            <div style="background: linear-gradient(135deg, #6366f1 0%, #8b5cf6 100%); padding: 30px; border-radius: 16px 16px 0 0; text-align: center;">
                <h1 style="color: white; margin: 0; font-size: 24px;">Interview Reminder</h1>
                <p style="color: rgba(255,255,255,0.9); margin: 10px 0 0 0;">Your interview is coming up!</p>
            </div>
            
            <div style="background: white; padding: 30px; border-radius: 0 0 16px 16px; box-shadow: 0 4px 6px rgba(0,0,0,0.1);">
                <p style="font-size: 16px; color: #334155; margin-bottom: 20px;">
                    Hi <strong>{candidate_name}</strong>,
                </p>
                
                <p style="font-size: 16px; color: #334155; margin-bottom: 25px;">
                    This is a friendly reminder that your <strong>{interview_type_display}</strong> at <strong>{company_name}</strong> is scheduled in <strong>{time_message}</strong>.
                </p>
                
                <div style="background: #f8fafc; border-radius: 12px; padding: 20px; margin-bottom: 25px;">
                    <h3 style="color: #1e293b; margin: 0 0 15px 0; font-size: 16px;">📋 Interview Details</h3>
                    <table style="width: 100%; border-collapse: collapse;">
                        <tr>
                            <td style="padding: 8px 0; color: #64748b; width: 120px;">Company:</td>
                            <td style="padding: 8px 0; font-weight: 600;">{company_name}</td>
                        </tr>
                        <tr>
                            <td style="padding: 8px 0; color: #64748b;">Position:</td>
                            <td style="padding: 8px 0;">{job_role}</td>
                        </tr>
                        <tr>
                            <td style="padding: 8px 0; color: #64748b;">Date:</td>
                            <td style="padding: 8px 0;">{interview_date}</td>
                        </tr>
                        <tr>
                            <td style="padding: 8px 0; color: #64748b;">Time:</td>
                            <td style="padding: 8px 0;">{interview_time}</td>
                        </tr>
                        <tr>
                            <td style="padding: 8px 0; color: #64748b;">Type:</td>
                            <td style="padding: 8px 0;">{interview_type_display}</td>
                        </tr>
                        {location_section}
                    </table>
                </div>
                
                <div style="background: #fef3c7; border-radius: 12px; padding: 20px; margin-bottom: 25px; border-left: 4px solid #f59e0b;">
                    <h3 style="color: #92400e; margin: 0 0 10px 0; font-size: 14px;">💡 Quick Preparation Tips</h3>
                    <ul style="color: #78350f; margin: 0; padding-left: 20px; font-size: 14px;">
                        <li>Review the job description and your resume</li>
                        <li>Prepare questions to ask the interviewer</li>
                        <li>Test your equipment if it's a video call</li>
                        <li>Have a copy of your resume ready</li>
                    </ul>
                </div>
                
                <p style="font-size: 14px; color: #64748b; margin-bottom: 0;">
                    Good luck with your interview! 🍀
                </p>
            </div>
            
            <div style="text-align: center; padding: 20px;">
                <p style="color: #94a3b8; font-size: 12px; margin: 0;">
                    Sent by HireFlow AI • Your AI-powered job search assistant
                </p>
            </div>
        </div>
    </body>
    </html>
    """
    
    return subject, html_body

async def send_interview_reminder(
    user_email: str,
    candidate_name: str,
    company_name: str,
    job_role: str,
    interview_date: str,
    interview_time: str,
    interview_type: str,
    location: str,
    meeting_link: str,
    reminder_type: str,
    event_id: str,
    user_id: str,
    job_application_id: str = None,
    retry: bool = True
) -> dict:
    """Send interview reminder email with retry logic"""
    
    if not RESEND_API_KEY:
        logger.warning("RESEND_API_KEY not configured, skipping email")
        return {"status": "skipped", "message": "Email service not configured"}
    
    subject, html_body = generate_interview_reminder_email(
        candidate_name=candidate_name,
        company_name=company_name,
        job_role=job_role,
        interview_date=interview_date,
        interview_time=interview_time,
        interview_type=interview_type,
        location=location or "",
        meeting_link=meeting_link or "",
        reminder_type=reminder_type
    )
    
    notification_id = str(uuid.uuid4())
    now = datetime.now(timezone.utc).isoformat()
    
    notification_doc = {
        "id": notification_id,
        "user_id": user_id,
        "event_id": event_id,
        "job_application_id": job_application_id,
        "reminder_type": reminder_type,
        "delivery_status": "pending",
        "error_message": None,
        "sent_timestamp": None,
        "recipient_email": user_email,
        "created_at": now
    }
    
    try:
        # Send email via Resend
        params = {
            "from": SENDER_EMAIL,
            "to": [user_email],
            "subject": subject,
            "html": html_body
        }
        
        email_response = resend.Emails.send(params)
        
        notification_doc["delivery_status"] = "sent"
        notification_doc["sent_timestamp"] = datetime.now(timezone.utc).isoformat()
        notification_doc["resend_id"] = email_response.get("id") if isinstance(email_response, dict) else str(email_response)
        
        await db.notification_logs.insert_one(notification_doc)
        logger.info(f"Email reminder sent: {reminder_type} for event {event_id} to {user_email}")
        
        return {"status": "sent", "notification_id": notification_id}
        
    except Exception as e:
        error_msg = str(e)
        logger.error(f"Email send failed: {error_msg}")
        
        if retry:
            # Retry once
            logger.info("Retrying email send...")
            try:
                email_response = resend.Emails.send(params)
                notification_doc["delivery_status"] = "sent"
                notification_doc["sent_timestamp"] = datetime.now(timezone.utc).isoformat()
                notification_doc["resend_id"] = email_response.get("id") if isinstance(email_response, dict) else str(email_response)
                await db.notification_logs.insert_one(notification_doc)
                logger.info(f"Email retry successful for event {event_id}")
                return {"status": "sent", "notification_id": notification_id}
            except Exception as retry_error:
                error_msg = f"Original: {error_msg}, Retry: {str(retry_error)}"
        
        notification_doc["delivery_status"] = "failed"
        notification_doc["error_message"] = error_msg
        await db.notification_logs.insert_one(notification_doc)
        
        return {"status": "failed", "error": error_msg, "notification_id": notification_id}

async def check_and_send_reminders():
    """Background job to check upcoming interviews and send reminders"""
    global scheduler_status
    start_time = datetime.now(timezone.utc)
    execution_log = []
    errors = []
    
    try:
        now = datetime.now(timezone.utc)
        
        log_msg = f"=== Reminder Check Started ==="
        logger.info(log_msg)
        execution_log.append({"time": now.isoformat(), "message": log_msg})
        
        log_msg = f"Server Time (UTC): {now.isoformat()}"
        logger.info(log_msg)
        execution_log.append({"time": now.isoformat(), "message": log_msg})
        
        # Calculate time windows for reminders
        # 24hr reminder: 23-25 hours before interview
        time_24hr_start = now + timedelta(hours=23)
        time_24hr_end = now + timedelta(hours=25)
        
        # 1hr reminder: 50-70 minutes before interview
        time_1hr_start = now + timedelta(minutes=50)
        time_1hr_end = now + timedelta(minutes=70)
        
        log_msg = f"24hr Window: {time_24hr_start.isoformat()} to {time_24hr_end.isoformat()}"
        logger.info(log_msg)
        execution_log.append({"time": now.isoformat(), "message": log_msg})
        
        log_msg = f"1hr Window: {time_1hr_start.isoformat()} to {time_1hr_end.isoformat()}"
        logger.info(log_msg)
        execution_log.append({"time": now.isoformat(), "message": log_msg})
        
        # Fetch ALL interview events with reminders enabled to debug
        all_events = await db.calendar_events.find({
            "reminders_enabled": {"$eq": True},
            "event_type": {"$in": ["interview", "phone_screen", "video_call"]}
        }, {"_id": 0}).to_list(100)
        
        log_msg = f"Total interview events with reminders enabled: {len(all_events)}"
        logger.info(log_msg)
        execution_log.append({"time": now.isoformat(), "message": log_msg})
        
        # Process events manually with proper datetime parsing
        events_24hr = []
        events_1hr = []
        
        for event in all_events:
            event_start_str = event.get("start_date", "")
            if not event_start_str:
                continue
                
            try:
                # Parse the event start date - handle various formats
                event_start = parse_event_datetime(event_start_str)
                
                # Calculate time until event
                time_until_event = event_start - now
                hours_until = time_until_event.total_seconds() / 3600
                
                logger.debug(f"Event '{event.get('title')}': starts at {event_start.isoformat()}, {hours_until:.2f} hours from now")
                
                # Check 24hr reminder (23-25 hours before)
                if 23 <= hours_until <= 25:
                    if not event.get("reminder_24hr_sent", False):
                        events_24hr.append(event)
                        log_msg = f"  -> Eligible for 24hr reminder: {event.get('title')}"
                        logger.info(log_msg)
                        execution_log.append({"time": now.isoformat(), "message": log_msg})
                
                # Check 1hr reminder (50-70 minutes before = 0.833-1.167 hours)
                if 0.833 <= hours_until <= 1.167:
                    if not event.get("reminder_1hr_sent", False):
                        events_1hr.append(event)
                        log_msg = f"  -> Eligible for 1hr reminder: {event.get('title')}"
                        logger.info(log_msg)
                        execution_log.append({"time": now.isoformat(), "message": log_msg})
                        
            except Exception as parse_error:
                error_msg = f"Failed to parse date for event {event.get('id')}: {parse_error}"
                logger.warning(error_msg)
                errors.append(error_msg)
                continue
        
        log_msg = f"Events eligible for 24hr reminder: {len(events_24hr)}"
        logger.info(log_msg)
        execution_log.append({"time": now.isoformat(), "message": log_msg})
        
        log_msg = f"Events eligible for 1hr reminder: {len(events_1hr)}"
        logger.info(log_msg)
        execution_log.append({"time": now.isoformat(), "message": log_msg})
        
        # Process reminders
        sent_24hr = 0
        sent_1hr = 0
        
        for event in events_24hr:
            log_msg = f"Sending 24hr reminder for: {event.get('title')}"
            logger.info(log_msg)
            execution_log.append({"time": datetime.now(timezone.utc).isoformat(), "message": log_msg})
            result = await process_reminder(event, "24hr")
            if result:
                sent_24hr += 1
        
        for event in events_1hr:
            log_msg = f"Sending 1hr reminder for: {event.get('title')}"
            logger.info(log_msg)
            execution_log.append({"time": datetime.now(timezone.utc).isoformat(), "message": log_msg})
            result = await process_reminder(event, "1hr")
            if result:
                sent_1hr += 1
        
        if events_24hr or events_1hr:
            log_msg = f"=== Reminder Check Complete: {len(events_24hr)} 24hr, {len(events_1hr)} 1hr reminders sent ==="
        else:
            log_msg = f"=== Reminder Check Complete: No reminders to send ==="
        logger.info(log_msg)
        execution_log.append({"time": datetime.now(timezone.utc).isoformat(), "message": log_msg})
        
        # Update scheduler status
        end_time = datetime.now(timezone.utc)
        scheduler_status = {
            "last_run_time": start_time.isoformat(),
            "last_run_duration_ms": int((end_time - start_time).total_seconds() * 1000),
            "total_events_checked": len(all_events),
            "eligible_24hr_count": len(events_24hr),
            "eligible_1hr_count": len(events_1hr),
            "reminders_sent_24hr": sent_24hr,
            "reminders_sent_1hr": sent_1hr,
            "last_errors": errors[-10:],  # Keep last 10 errors
            "execution_logs": execution_log[-50:]  # Keep last 50 log entries
        }
            
    except Exception as e:
        error_msg = f"Reminder check failed: {e}"
        logger.error(error_msg, exc_info=True)
        errors.append(error_msg)
        scheduler_status["last_errors"] = errors[-10:]
        scheduler_status["last_run_time"] = start_time.isoformat()

def parse_event_datetime(date_str: str) -> datetime:
    """Parse event datetime string to timezone-aware datetime"""
    if not date_str:
        raise ValueError("Empty date string")
    
    # Handle ISO format with Z suffix
    if date_str.endswith('Z'):
        date_str = date_str[:-1] + '+00:00'
    
    # Try parsing with timezone info
    try:
        dt = datetime.fromisoformat(date_str)
        # If no timezone, assume UTC
        if dt.tzinfo is None:
            dt = dt.replace(tzinfo=timezone.utc)
        return dt
    except ValueError:
        pass
    
    # Try common formats without timezone
    for fmt in [
        "%Y-%m-%dT%H:%M:%S",
        "%Y-%m-%dT%H:%M",
        "%Y-%m-%d %H:%M:%S",
        "%Y-%m-%d %H:%M"
    ]:
        try:
            dt = datetime.strptime(date_str, fmt)
            return dt.replace(tzinfo=timezone.utc)
        except ValueError:
            continue
    
    raise ValueError(f"Unable to parse date: {date_str}")

async def process_reminder(event: dict, reminder_type: str) -> bool:
    """Process a single reminder. Returns True if sent successfully."""
    try:
        event_id = event.get("id")
        
        # CRITICAL: Re-fetch the event from database to get the latest reminders_enabled status
        # This prevents sending reminders if user disabled them after the initial query
        fresh_event = await db.calendar_events.find_one({"id": event_id}, {"_id": 0})
        if not fresh_event:
            logger.warning(f"Event {event_id} no longer exists, skipping reminder")
            return False
        
        # Final validation: Check if reminders are still enabled
        if not fresh_event.get("reminders_enabled", False):
            logger.info(f"Reminders disabled for event {event_id}, skipping {reminder_type} reminder")
            return False
        
        # Check if this specific reminder was already sent (in case of race condition)
        reminder_sent_field = f"reminder_{reminder_type}_sent"
        if fresh_event.get(reminder_sent_field, False):
            logger.info(f"Reminder {reminder_type} already sent for event {event_id}, skipping")
            return False
        
        # Get user info
        user = await db.users.find_one({"id": fresh_event["user_id"]}, {"_id": 0})
        if not user:
            logger.warning(f"User not found for event {event_id}")
            return False
        
        # Get job application info if linked
        company_name = "Company"
        job_role = fresh_event.get("title", "Interview")
        
        if fresh_event.get("job_application_id"):
            app = await db.applications.find_one({"id": fresh_event["job_application_id"]}, {"_id": 0})
            if app:
                company_name = app.get("company", company_name)
                job_role = app.get("position", job_role)
        
        # Parse interview date/time
        start_date = datetime.fromisoformat(fresh_event["start_date"].replace("Z", "+00:00"))
        interview_date = start_date.strftime("%B %d, %Y")
        interview_time = start_date.strftime("%I:%M %p")
        
        # Send reminder
        result = await send_interview_reminder(
            user_email=user["email"],
            candidate_name=user["name"],
            company_name=company_name,
            job_role=job_role,
            interview_date=interview_date,
            interview_time=interview_time,
            interview_type=fresh_event.get("interview_type", fresh_event.get("event_type", "interview")),
            location=fresh_event.get("location"),
            meeting_link=fresh_event.get("meeting_link"),
            reminder_type=reminder_type,
            event_id=event_id,
            user_id=fresh_event["user_id"],
            job_application_id=fresh_event.get("job_application_id")
        )
        
        # Mark reminder as sent
        if result["status"] == "sent":
            await db.calendar_events.update_one(
                {"id": event_id},
                {"$set": {reminder_sent_field: True}}
            )
            return True
        return False
            
    except Exception as e:
        logger.error(f"Failed to process reminder for event {event.get('id')}: {e}")
        return False

# ===================== AUTH ROUTES =====================

@api_router.post("/auth/register", response_model=TokenResponse)
async def register(user_data: UserCreate):
    existing = await db.users.find_one({"email": user_data.email})
    if existing:
        raise HTTPException(status_code=400, detail="Email already registered")
    
    user_id = str(uuid.uuid4())
    now = datetime.now(timezone.utc).isoformat()
    
    user_doc = {
        "id": user_id,
        "email": user_data.email,
        "name": user_data.name,
        "password": hash_password(user_data.password),
        "created_at": now
    }
    
    await db.users.insert_one(user_doc)
    token = create_token(user_id)
    
    return TokenResponse(
        token=token,
        user=UserResponse(id=user_id, email=user_data.email, name=user_data.name, created_at=now)
    )

@api_router.post("/auth/login", response_model=TokenResponse)
async def login(credentials: UserLogin):
    user = await db.users.find_one({"email": credentials.email}, {"_id": 0})
    if not user or not verify_password(credentials.password, user["password"]):
        raise HTTPException(status_code=401, detail="Invalid credentials")
    
    token = create_token(user["id"])
    return TokenResponse(
        token=token,
        user=UserResponse(id=user["id"], email=user["email"], name=user["name"], created_at=user["created_at"])
    )

@api_router.get("/auth/me", response_model=UserResponse)
async def get_me(current_user: dict = Depends(get_current_user)):
    return UserResponse(**current_user)

# ===================== RESUME ROUTES =====================

@api_router.post("/resumes", response_model=ResumeResponse)
async def create_resume(resume_data: ResumeCreate, current_user: dict = Depends(get_current_user)):
    resume_id = str(uuid.uuid4())
    now = datetime.now(timezone.utc).isoformat()
    
    resume_doc = {
        "id": resume_id,
        "user_id": current_user["id"],
        "title": resume_data.title,
        "content": resume_data.content,
        "file_name": None,
        "file_type": None,
        "extraction_method": "manual",
        "extraction_status": "success",
        "ocr_used": False,
        "analysis": None,
        "score": None,
        "created_at": now,
        "updated_at": now
    }
    
    await db.resumes.insert_one(resume_doc)
    return ResumeResponse(**{k: v for k, v in resume_doc.items() if k != "_id"})

@api_router.post("/resumes/upload", response_model=ResumeResponse)
async def upload_resume(
    file: UploadFile = File(...),
    title: str = Form(...),
    current_user: dict = Depends(get_current_user)
):
    """Upload a PDF or DOCX resume file and extract text content with OCR fallback"""
    # Validate file type
    allowed_types = {
        "application/pdf": "pdf",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
        "application/msword": "doc",
        "image/png": "png",
        "image/jpeg": "jpg",
        "image/jpg": "jpg"
    }
    
    content_type = file.content_type
    if content_type not in allowed_types:
        # Also check by file extension
        file_ext = file.filename.lower().split(".")[-1] if file.filename else ""
        if file_ext not in ["pdf", "docx", "doc", "png", "jpg", "jpeg"]:
            raise HTTPException(
                status_code=400, 
                detail="Invalid file type. Please upload a PDF, DOCX, or image file."
            )
        file_type = file_ext
    else:
        file_type = allowed_types[content_type]
    
    # Read file content
    file_content = await file.read()
    
    # Check file size (max 10MB)
    if len(file_content) > 10 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="File size exceeds 10MB limit")
    
    # Extract text based on file type with OCR fallback
    extraction_result = None
    
    if file_type == "pdf":
        logger.info(f"Processing PDF file: {file.filename}")
        extraction_result = extract_text_from_pdf(file_content)
    elif file_type in ["docx", "doc"]:
        logger.info(f"Processing DOCX file: {file.filename}")
        extraction_result = extract_text_from_docx(file_content)
    elif file_type in ["png", "jpg", "jpeg"]:
        logger.info(f"Processing image file with OCR: {file.filename}")
        ocr_text = extract_text_from_image_ocr(file_content)
        if is_text_readable(ocr_text):
            extraction_result = ExtractionResult(
                text=ocr_text,
                method="ocr",
                status="success",
                ocr_used=True
            )
        elif ocr_text:
            extraction_result = ExtractionResult(
                text=ocr_text,
                method="ocr",
                status="partial",
                ocr_used=True
            )
        else:
            extraction_result = ExtractionResult(
                text="",
                method="failed",
                status="failed",
                ocr_used=True
            )
    else:
        raise HTTPException(status_code=400, detail="Unsupported file format")
    
    # Handle extraction failures
    if extraction_result.status == "failed":
        raise HTTPException(
            status_code=400, 
            detail="Could not extract text from the file. Please upload an ATS-friendly resume (searchable PDF or DOCX) with clear, readable text."
        )
    
    # Warn about partial extraction but still process
    if extraction_result.status == "partial" and len(extraction_result.text) < 50:
        raise HTTPException(
            status_code=400, 
            detail="Extracted text is too short. Please upload a resume with more content, preferably in DOCX or searchable PDF format."
        )
    
    # Create resume record with extraction metadata
    resume_id = str(uuid.uuid4())
    now = datetime.now(timezone.utc).isoformat()
    
    resume_doc = {
        "id": resume_id,
        "user_id": current_user["id"],
        "title": title,
        "content": extraction_result.text,
        "file_name": file.filename,
        "file_type": file_type,
        "extraction_method": extraction_result.method,
        "extraction_status": extraction_result.status,
        "ocr_used": extraction_result.ocr_used,
        "analysis": None,
        "score": None,
        "created_at": now,
        "updated_at": now
    }
    
    await db.resumes.insert_one(resume_doc)
    
    log_msg = f"Resume uploaded: {file.filename} | Method: {extraction_result.method} | OCR: {extraction_result.ocr_used} | Status: {extraction_result.status} | Chars: {len(extraction_result.text)}"
    logger.info(log_msg)
    
    return ResumeResponse(**{k: v for k, v in resume_doc.items() if k != "_id"})

@api_router.get("/resumes", response_model=List[ResumeResponse])
async def get_resumes(current_user: dict = Depends(get_current_user)):
    resumes = await db.resumes.find({"user_id": current_user["id"]}, {"_id": 0}).to_list(100)
    return [ResumeResponse(**r) for r in resumes]

@api_router.get("/resumes/{resume_id}", response_model=ResumeResponse)
async def get_resume(resume_id: str, current_user: dict = Depends(get_current_user)):
    resume = await db.resumes.find_one({"id": resume_id, "user_id": current_user["id"]}, {"_id": 0})
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")
    return ResumeResponse(**resume)

@api_router.post("/resumes/{resume_id}/analyze", response_model=ResumeResponse)
async def analyze_resume(resume_id: str, current_user: dict = Depends(get_current_user)):
    resume = await db.resumes.find_one({"id": resume_id, "user_id": current_user["id"]}, {"_id": 0})
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")
    
    system_msg = """You are an expert resume analyst. Analyze the resume and provide:
1. An overall score from 0-100
2. Key strengths (list 3-5)
3. Areas for improvement (list 3-5)
4. Industry keywords present
5. Formatting suggestions

Respond in JSON format:
{
    "score": <number>,
    "strengths": ["..."],
    "improvements": ["..."],
    "keywords": ["..."],
    "formatting_tips": ["..."],
    "summary": "..."
}"""
    
    response = await get_llm_response(system_msg, f"Analyze this resume:\n\n{resume['content']}")
    
    try:
        import json
        # Clean the response - remove markdown code blocks if present
        clean_response = response.strip()
        if clean_response.startswith("```json"):
            clean_response = clean_response[7:]
        if clean_response.startswith("```"):
            clean_response = clean_response[3:]
        if clean_response.endswith("```"):
            clean_response = clean_response[:-3]
        analysis = json.loads(clean_response.strip())
        score = analysis.get("score", 70)
    except:
        analysis = {"summary": response, "score": 70}
        score = 70
    
    now = datetime.now(timezone.utc).isoformat()
    await db.resumes.update_one(
        {"id": resume_id},
        {"$set": {"analysis": analysis, "score": score, "updated_at": now}}
    )
    
    resume["analysis"] = analysis
    resume["score"] = score
    resume["updated_at"] = now
    return ResumeResponse(**resume)

@api_router.delete("/resumes/{resume_id}")
async def delete_resume(resume_id: str, current_user: dict = Depends(get_current_user)):
    result = await db.resumes.delete_one({"id": resume_id, "user_id": current_user["id"]})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Resume not found")
    return {"message": "Resume deleted"}

# ===================== JOB APPLICATION ROUTES =====================

@api_router.post("/applications", response_model=JobApplicationResponse)
async def create_application(app_data: JobApplicationCreate, current_user: dict = Depends(get_current_user)):
    app_id = str(uuid.uuid4())
    now = datetime.now(timezone.utc).isoformat()
    applied_date = app_data.applied_date or now
    
    app_doc = {
        "id": app_id,
        "user_id": current_user["id"],
        "company": app_data.company,
        "position": app_data.position,
        "job_url": app_data.job_url,
        "job_description": app_data.job_description,
        "status": app_data.status,
        "notes": app_data.notes,
        "applied_date": applied_date,
        "created_at": now,
        "updated_at": now
    }
    
    await db.applications.insert_one(app_doc)
    return JobApplicationResponse(**{k: v for k, v in app_doc.items() if k != "_id"})

@api_router.get("/applications", response_model=List[JobApplicationResponse])
async def get_applications(status: Optional[str] = None, current_user: dict = Depends(get_current_user)):
    query = {"user_id": current_user["id"]}
    if status:
        query["status"] = status
    applications = await db.applications.find(query, {"_id": 0}).sort("created_at", -1).to_list(500)
    return [JobApplicationResponse(**a) for a in applications]

@api_router.get("/applications/{app_id}", response_model=JobApplicationResponse)
async def get_application(app_id: str, current_user: dict = Depends(get_current_user)):
    application = await db.applications.find_one({"id": app_id, "user_id": current_user["id"]}, {"_id": 0})
    if not application:
        raise HTTPException(status_code=404, detail="Application not found")
    return JobApplicationResponse(**application)

@api_router.put("/applications/{app_id}", response_model=JobApplicationResponse)
async def update_application(app_id: str, app_data: JobApplicationUpdate, current_user: dict = Depends(get_current_user)):
    update_data = {k: v for k, v in app_data.model_dump().items() if v is not None}
    if not update_data:
        raise HTTPException(status_code=400, detail="No update data provided")
    
    update_data["updated_at"] = datetime.now(timezone.utc).isoformat()
    
    result = await db.applications.update_one(
        {"id": app_id, "user_id": current_user["id"]},
        {"$set": update_data}
    )
    
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Application not found")
    
    updated = await db.applications.find_one({"id": app_id}, {"_id": 0})
    return JobApplicationResponse(**updated)

@api_router.delete("/applications/{app_id}")
async def delete_application(app_id: str, current_user: dict = Depends(get_current_user)):
    result = await db.applications.delete_one({"id": app_id, "user_id": current_user["id"]})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Application not found")
    return {"message": "Application deleted"}

# ===================== COVER LETTER ROUTES (Phase 1) =====================

@api_router.post("/cover-letter/generate", response_model=CoverLetterFullResponse)
async def generate_cover_letter_v2(request: CoverLetterGenerateRequest, current_user: dict = Depends(get_current_user)):
    """
    Phase 1 AI Cover Letter Generator
    - Fetches resume content and job application details
    - Generates ATS-friendly cover letter using AI
    - Stores with full metadata
    """
    # Fetch resume
    resume = await db.resumes.find_one({"id": request.resume_id, "user_id": current_user["id"]}, {"_id": 0})
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")
    
    # Fetch job application
    job_app = await db.applications.find_one({"id": request.job_application_id, "user_id": current_user["id"]}, {"_id": 0})
    if not job_app:
        raise HTTPException(status_code=404, detail="Job application not found")
    
    resume_content = resume.get("content", "")
    company_name = job_app.get("company", "")
    position = job_app.get("position", "")
    job_description = job_app.get("job_description", "")
    
    # AI prompt for ATS-friendly cover letter
    system_msg = """You are an expert cover letter writer specializing in ATS-friendly, professional cover letters.

REQUIREMENTS:
- Write a professional, tailored cover letter
- Use a 3-4 paragraph structure:
  1. Opening: Express interest in the specific role and company
  2. Skills Alignment: Match your experience to the job requirements
  3. Key Achievements: Highlight 2-3 relevant accomplishments from the resume
  4. Closing: Express enthusiasm and gratitude, include call to action
- Keep between 250-400 words
- Use professional tone throughout
- IMPORTANT: Only reference experience, skills, and achievements that are actually in the resume - DO NOT hallucinate or invent experiences
- Format for ATS compatibility (clean paragraphs, no special formatting)
- Include a professional greeting and sign-off

Output ONLY the cover letter text, ready to use."""

    user_msg = f"""Generate a cover letter for this job application:

COMPANY: {company_name}
POSITION: {position}
JOB DESCRIPTION:
{job_description if job_description else "Not provided - use the position title to tailor the letter"}

CANDIDATE'S RESUME:
{resume_content}

{f"ADDITIONAL NOTES FROM CANDIDATE: {request.customization_notes}" if request.customization_notes else ""}

Write the cover letter now:"""

    try:
        content = await get_llm_response(system_msg, user_msg)
    except Exception as e:
        logger.error(f"AI generation failed: {e}")
        raise HTTPException(status_code=500, detail="Failed to generate cover letter. Please try again.")
    
    # Calculate word count
    word_count = len(content.split())
    
    letter_id = str(uuid.uuid4())
    now = datetime.now(timezone.utc).isoformat()
    
    letter_doc = {
        "id": letter_id,
        "user_id": current_user["id"],
        "resume_id": request.resume_id,
        "job_application_id": request.job_application_id,
        "title": f"Cover Letter - {position} at {company_name}",
        "content": content,
        "word_count": word_count,
        "company_name": company_name,
        "position": position,
        "customization_notes": request.customization_notes,
        "created_at": now,
        "updated_at": now
    }
    
    await db.cover_letters_v2.insert_one(letter_doc)
    return CoverLetterFullResponse(**{k: v for k, v in letter_doc.items() if k not in ["_id", "customization_notes"]})

@api_router.get("/cover-letter", response_model=List[CoverLetterFullResponse])
async def get_all_cover_letters(current_user: dict = Depends(get_current_user)):
    """Get all cover letters for the logged-in user"""
    letters = await db.cover_letters_v2.find(
        {"user_id": current_user["id"]}, 
        {"_id": 0, "customization_notes": 0}
    ).sort("created_at", -1).to_list(100)
    return [CoverLetterFullResponse(**l) for l in letters]

@api_router.get("/cover-letter/{letter_id}", response_model=CoverLetterFullResponse)
async def get_cover_letter(letter_id: str, current_user: dict = Depends(get_current_user)):
    """Get a single cover letter by ID"""
    letter = await db.cover_letters_v2.find_one(
        {"id": letter_id, "user_id": current_user["id"]}, 
        {"_id": 0, "customization_notes": 0}
    )
    if not letter:
        raise HTTPException(status_code=404, detail="Cover letter not found")
    return CoverLetterFullResponse(**letter)

@api_router.put("/cover-letter/{letter_id}", response_model=CoverLetterFullResponse)
async def update_cover_letter(letter_id: str, update_data: CoverLetterUpdateRequest, current_user: dict = Depends(get_current_user)):
    """Update a cover letter (content or title)"""
    update_fields = {}
    
    if update_data.content is not None:
        update_fields["content"] = update_data.content
        update_fields["word_count"] = len(update_data.content.split())
    
    if update_data.title is not None:
        update_fields["title"] = update_data.title
    
    if not update_fields:
        raise HTTPException(status_code=400, detail="No update data provided")
    
    update_fields["updated_at"] = datetime.now(timezone.utc).isoformat()
    
    result = await db.cover_letters_v2.update_one(
        {"id": letter_id, "user_id": current_user["id"]},
        {"$set": update_fields}
    )
    
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Cover letter not found")
    
    updated = await db.cover_letters_v2.find_one({"id": letter_id}, {"_id": 0, "customization_notes": 0})
    return CoverLetterFullResponse(**updated)

@api_router.delete("/cover-letter/{letter_id}")
async def delete_cover_letter_v2(letter_id: str, current_user: dict = Depends(get_current_user)):
    """Delete a cover letter"""
    result = await db.cover_letters_v2.delete_one({"id": letter_id, "user_id": current_user["id"]})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Cover letter not found")
    return {"message": "Cover letter deleted"}

@api_router.get("/cover-letter/{letter_id}/pdf")
async def download_cover_letter_pdf(letter_id: str, current_user: dict = Depends(get_current_user)):
    """Generate and download cover letter as PDF"""
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from fastapi.responses import StreamingResponse
    
    # Fetch cover letter
    cover_letter = await db.cover_letters_v2.find_one(
        {"id": letter_id, "user_id": current_user["id"]}, 
        {"_id": 0}
    )
    if not cover_letter:
        raise HTTPException(status_code=404, detail="Cover letter not found")
    
    # Get user info for header
    user = await db.users.find_one({"id": current_user["id"]}, {"_id": 0})
    
    # Create PDF in memory
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        rightMargin=inch,
        leftMargin=inch,
        topMargin=inch,
        bottomMargin=inch
    )
    
    # Styles
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'Title',
        parent=styles['Heading1'],
        fontSize=14,
        spaceAfter=12
    )
    body_style = ParagraphStyle(
        'Body',
        parent=styles['Normal'],
        fontSize=11,
        leading=14,
        spaceAfter=12
    )
    
    # Build PDF content
    story = []
    
    # Header with user name
    if user:
        story.append(Paragraph(user.get("name", ""), title_style))
        story.append(Paragraph(user.get("email", ""), body_style))
        story.append(Spacer(1, 0.3 * inch))
    
    # Date
    story.append(Paragraph(datetime.now().strftime("%B %d, %Y"), body_style))
    story.append(Spacer(1, 0.2 * inch))
    
    # Company info
    story.append(Paragraph(f"RE: {cover_letter['position']} Position", body_style))
    story.append(Paragraph(cover_letter['company_name'], body_style))
    story.append(Spacer(1, 0.3 * inch))
    
    # Cover letter content - split by paragraphs
    content = cover_letter['content']
    paragraphs = content.split('\n\n')
    for para in paragraphs:
        if para.strip():
            # Clean up any single newlines within paragraphs
            cleaned_para = para.replace('\n', ' ').strip()
            story.append(Paragraph(cleaned_para, body_style))
    
    # Build PDF
    doc.build(story)
    buffer.seek(0)
    
    # Generate filename
    safe_company = "".join(c for c in cover_letter['company_name'] if c.isalnum() or c in (' ', '-', '_')).strip()
    safe_position = "".join(c for c in cover_letter['position'] if c.isalnum() or c in (' ', '-', '_')).strip()
    filename = f"Cover_Letter_{safe_company}_{safe_position}.pdf"
    
    return StreamingResponse(
        buffer,
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )

# Legacy cover letter routes (keeping for backward compatibility)
@api_router.post("/cover-letters/generate", response_model=CoverLetterResponse)
async def generate_cover_letter(request: CoverLetterRequest, current_user: dict = Depends(get_current_user)):
    resume_content = ""
    if request.resume_id:
        resume = await db.resumes.find_one({"id": request.resume_id, "user_id": current_user["id"]}, {"_id": 0})
        if resume:
            resume_content = resume.get("content", "")
    
    system_msg = """You are an expert cover letter writer. Write a professional, compelling cover letter that:
1. Is tailored to the specific job and company
2. Highlights relevant experience and skills
3. Shows enthusiasm for the role
4. Is concise (300-400 words)
5. Has a professional tone but shows personality

Format it properly with greeting, body paragraphs, and professional closing."""

    user_msg = f"""Write a cover letter for:
Company: {request.company_name}
Position: {request.position}
Job Description: {request.job_description}

{"My Resume:" + resume_content if resume_content else ""}"""

    content = await get_llm_response(system_msg, user_msg)
    
    letter_id = str(uuid.uuid4())
    now = datetime.now(timezone.utc).isoformat()
    
    letter_doc = {
        "id": letter_id,
        "user_id": current_user["id"],
        "company_name": request.company_name,
        "position": request.position,
        "content": content,
        "created_at": now
    }
    
    await db.cover_letters.insert_one(letter_doc)
    return CoverLetterResponse(**{k: v for k, v in letter_doc.items() if k != "_id"})

@api_router.get("/cover-letters", response_model=List[CoverLetterResponse])
async def get_cover_letters(current_user: dict = Depends(get_current_user)):
    letters = await db.cover_letters.find({"user_id": current_user["id"]}, {"_id": 0}).sort("created_at", -1).to_list(100)
    return [CoverLetterResponse(**l) for l in letters]

@api_router.delete("/cover-letters/{letter_id}")
async def delete_cover_letter(letter_id: str, current_user: dict = Depends(get_current_user)):
    result = await db.cover_letters.delete_one({"id": letter_id, "user_id": current_user["id"]})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Cover letter not found")
    return {"message": "Cover letter deleted"}

# ===================== JOB MATCHING ROUTES =====================

@api_router.post("/match/analyze", response_model=JobMatchResponse)
async def analyze_resume_job_match(request: MatchRequest, current_user: dict = Depends(get_current_user)):
    """Enhanced resume-to-job matching with comprehensive analysis stored in database"""
    resume = await db.resumes.find_one({"id": request.resume_id, "user_id": current_user["id"]}, {"_id": 0})
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")
    
    system_msg = """You are an expert career advisor and resume analyst. Analyze how well the resume matches the job description.

Provide a comprehensive, structured analysis including:
1. Overall match score (0-100) based on skills, experience, and qualifications alignment
2. Skill match analysis - which required skills are present vs missing
3. Experience match - how well the experience level and type matches
4. Missing skills - specific skills from the job that are not in the resume
5. Weak areas - skills mentioned but not strongly demonstrated
6. Strengths - areas where the candidate excels for this role
7. Actionable suggestions - specific, role-targeted improvements
8. Keyword analysis - important keywords from job description found/missing in resume

Respond ONLY in this exact JSON format:
{
    "match_score": <number 0-100>,
    "skill_match": {
        "matched_skills": ["skill1", "skill2"],
        "partial_match": ["skill3"],
        "missing_skills": ["skill4", "skill5"]
    },
    "experience_match": {
        "score": <number 0-100>,
        "analysis": "brief explanation of experience alignment"
    },
    "missing_skills": ["specific skill 1", "specific skill 2"],
    "weak_areas": ["area needing improvement 1", "area 2"],
    "strengths": ["strength 1", "strength 2", "strength 3"],
    "suggestions": [
        "Specific actionable suggestion 1",
        "Specific actionable suggestion 2",
        "Specific actionable suggestion 3"
    ],
    "keyword_analysis": {
        "found": ["keyword1", "keyword2"],
        "missing": ["keyword3", "keyword4"],
        "recommendation": "Add these keywords naturally to your resume"
    },
    "summary": "2-3 sentence summary of the match and key recommendations"
}"""

    job_context = f"Job Title: {request.job_title}\nCompany: {request.company_name}\n" if request.job_title else ""
    
    user_msg = f"""Analyze this resume against the job description:

{job_context}
RESUME CONTENT:
{resume['content']}

JOB DESCRIPTION:
{request.job_description}

Provide detailed, actionable analysis."""

    response = await get_llm_response(system_msg, user_msg)
    
    try:
        import json
        clean_response = response.strip()
        if clean_response.startswith("```json"):
            clean_response = clean_response[7:]
        if clean_response.startswith("```"):
            clean_response = clean_response[3:]
        if clean_response.endswith("```"):
            clean_response = clean_response[:-3]
        analysis_data = json.loads(clean_response.strip())
        
        # Ensure all required fields exist with defaults
        analysis = MatchAnalysis(
            match_score=analysis_data.get("match_score", 50),
            skill_match=analysis_data.get("skill_match", {"matched_skills": [], "partial_match": [], "missing_skills": []}),
            experience_match=analysis_data.get("experience_match", {"score": 50, "analysis": "Analysis not available"}),
            missing_skills=analysis_data.get("missing_skills", []),
            weak_areas=analysis_data.get("weak_areas", []),
            strengths=analysis_data.get("strengths", []),
            suggestions=analysis_data.get("suggestions", []),
            keyword_analysis=analysis_data.get("keyword_analysis", {"found": [], "missing": [], "recommendation": ""}),
            summary=analysis_data.get("summary", "Analysis completed")
        )
    except Exception as e:
        logger.error(f"Failed to parse LLM response: {e}")
        analysis = MatchAnalysis(
            match_score=50,
            skill_match={"matched_skills": [], "partial_match": [], "missing_skills": []},
            experience_match={"score": 50, "analysis": "Unable to analyze experience match"},
            missing_skills=["Unable to determine - please try again"],
            weak_areas=[],
            strengths=["Resume content detected"],
            suggestions=["Please try the analysis again for detailed results"],
            keyword_analysis={"found": [], "missing": [], "recommendation": "Retry analysis"},
            summary="Analysis encountered an error. Please try again."
        )
    
    # Store the analysis in database
    match_id = str(uuid.uuid4())
    now = datetime.now(timezone.utc).isoformat()
    
    match_doc = {
        "id": match_id,
        "user_id": current_user["id"],
        "resume_id": request.resume_id,
        "job_title": request.job_title,
        "company_name": request.company_name,
        "job_description": request.job_description,
        "analysis": analysis.model_dump(),
        "created_at": now
    }
    
    await db.job_matches.insert_one(match_doc)
    logger.info(f"Job match analysis saved: {match_id} for user {current_user['id']}")
    
    return JobMatchResponse(**{k: v for k, v in match_doc.items() if k != "_id"})

@api_router.get("/match/history", response_model=List[JobMatchResponse])
async def get_match_history(resume_id: Optional[str] = None, current_user: dict = Depends(get_current_user)):
    """Get all job match analyses for the user, optionally filtered by resume"""
    query = {"user_id": current_user["id"]}
    if resume_id:
        query["resume_id"] = resume_id
    
    matches = await db.job_matches.find(query, {"_id": 0}).sort("created_at", -1).to_list(100)
    return [JobMatchResponse(**m) for m in matches]

@api_router.get("/match/{match_id}", response_model=JobMatchResponse)
async def get_match_analysis(match_id: str, current_user: dict = Depends(get_current_user)):
    """Get a specific job match analysis"""
    match = await db.job_matches.find_one({"id": match_id, "user_id": current_user["id"]}, {"_id": 0})
    if not match:
        raise HTTPException(status_code=404, detail="Match analysis not found")
    return JobMatchResponse(**match)

@api_router.delete("/match/{match_id}")
async def delete_match_analysis(match_id: str, current_user: dict = Depends(get_current_user)):
    """Delete a job match analysis"""
    result = await db.job_matches.delete_one({"id": match_id, "user_id": current_user["id"]})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Match analysis not found")
    return {"message": "Match analysis deleted"}

@api_router.post("/match", response_model=MatchResponse)
async def match_resume_to_job(request: MatchRequest, current_user: dict = Depends(get_current_user)):
    """Legacy simple match endpoint for backward compatibility"""
    resume = await db.resumes.find_one({"id": request.resume_id, "user_id": current_user["id"]}, {"_id": 0})
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")
    
    system_msg = """You are an expert job matching analyst. Compare the resume to the job description and provide:
1. A match score (0-100)
2. Key strengths that align with the job
3. Gaps or missing qualifications
4. Suggestions to improve the match

Respond in JSON format:
{
    "match_score": <number>,
    "strengths": ["..."],
    "gaps": ["..."],
    "suggestions": ["..."]
}"""

    user_msg = f"""Compare this resume to the job description:

RESUME:
{resume['content']}

JOB DESCRIPTION:
{request.job_description}"""

    response = await get_llm_response(system_msg, user_msg)
    
    try:
        import json
        clean_response = response.strip()
        if clean_response.startswith("```json"):
            clean_response = clean_response[7:]
        if clean_response.startswith("```"):
            clean_response = clean_response[3:]
        if clean_response.endswith("```"):
            clean_response = clean_response[:-3]
        data = json.loads(clean_response.strip())
        return MatchResponse(**data)
    except:
        return MatchResponse(
            match_score=70,
            strengths=["Unable to parse detailed analysis"],
            gaps=[],
            suggestions=["Please try again"]
        )

# ===================== CALENDAR ROUTES =====================

@api_router.post("/calendar", response_model=CalendarEventResponse)
async def create_event(event_data: CalendarEventCreate, current_user: dict = Depends(get_current_user)):
    event_id = str(uuid.uuid4())
    now = datetime.now(timezone.utc).isoformat()
    
    event_doc = {
        "id": event_id,
        "user_id": current_user["id"],
        "title": event_data.title,
        "description": event_data.description,
        "event_type": event_data.event_type,
        "interview_type": event_data.interview_type,
        "start_date": event_data.start_date,
        "end_date": event_data.end_date,
        "job_application_id": event_data.job_application_id,
        "location": event_data.location,
        "meeting_link": event_data.meeting_link,
        "notes": event_data.notes,
        "reminders_enabled": event_data.reminders_enabled,
        "reminder_24hr_sent": False,
        "reminder_1hr_sent": False,
        "created_at": now
    }
    
    await db.calendar_events.insert_one(event_doc)
    
    # Log if reminders are enabled for interview-type events
    if event_data.reminders_enabled and event_data.event_type in ["interview", "phone_screen", "video_call"]:
        logger.info(f"Interview event created with reminders enabled: {event_id} at {event_data.start_date}")
    
    return CalendarEventResponse(**{k: v for k, v in event_doc.items() if k != "_id"})

@api_router.get("/calendar", response_model=List[CalendarEventResponse])
async def get_events(month: Optional[str] = None, current_user: dict = Depends(get_current_user)):
    query = {"user_id": current_user["id"]}
    events = await db.calendar_events.find(query, {"_id": 0}).sort("start_date", 1).to_list(500)
    return [CalendarEventResponse(**e) for e in events]

@api_router.put("/calendar/{event_id}", response_model=CalendarEventResponse)
async def update_event(event_id: str, event_data: CalendarEventUpdate, current_user: dict = Depends(get_current_user)):
    # Build update data - handle boolean False values correctly (don't filter them out)
    update_data = {}
    for k, v in event_data.model_dump().items():
        # Include field if it's not None, OR if it's a boolean (False is a valid value)
        if v is not None or isinstance(v, bool):
            if v is not None:  # Only add non-None values
                update_data[k] = v
    
    if not update_data:
        raise HTTPException(status_code=400, detail="No update data provided")
    
    # If start_date is changed, reset reminder flags so they can be sent again
    if "start_date" in update_data:
        update_data["reminder_24hr_sent"] = False
        update_data["reminder_1hr_sent"] = False
    
    # CRITICAL: If reminders are being disabled, mark all reminders as "sent" 
    # to prevent the background scheduler from sending them
    if "reminders_enabled" in update_data and update_data["reminders_enabled"] is False:
        update_data["reminder_24hr_sent"] = True  # Prevents future 24hr reminder
        update_data["reminder_1hr_sent"] = True   # Prevents future 1hr reminder
        logger.info(f"Reminders disabled for event {event_id}, marking reminders as complete to prevent sending")
    
    result = await db.calendar_events.update_one(
        {"id": event_id, "user_id": current_user["id"]},
        {"$set": update_data}
    )
    
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Event not found")
    
    updated = await db.calendar_events.find_one({"id": event_id}, {"_id": 0})
    return CalendarEventResponse(**updated)

@api_router.delete("/calendar/{event_id}")
async def delete_event(event_id: str, current_user: dict = Depends(get_current_user)):
    result = await db.calendar_events.delete_one({"id": event_id, "user_id": current_user["id"]})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Event not found")
    return {"message": "Event deleted"}

@api_router.post("/calendar/{event_id}/test-reminder")
async def send_test_reminder(event_id: str, current_user: dict = Depends(get_current_user)):
    """Send a test reminder email for an event (for testing purposes)"""
    event = await db.calendar_events.find_one({"id": event_id, "user_id": current_user["id"]}, {"_id": 0})
    if not event:
        raise HTTPException(status_code=404, detail="Event not found")
    
    if not RESEND_API_KEY:
        raise HTTPException(status_code=503, detail="Email service not configured")
    
    # Get user info
    user = await db.users.find_one({"id": current_user["id"]}, {"_id": 0})
    
    # Get job application info if linked
    company_name = "Your Company"
    job_role = event.get("title", "Interview")
    
    if event.get("job_application_id"):
        app = await db.applications.find_one({"id": event["job_application_id"]}, {"_id": 0})
        if app:
            company_name = app.get("company", company_name)
            job_role = app.get("position", job_role)
    
    # Parse interview date/time
    start_date = datetime.fromisoformat(event["start_date"].replace("Z", "+00:00"))
    interview_date = start_date.strftime("%B %d, %Y")
    interview_time = start_date.strftime("%I:%M %p")
    
    # Send test reminder
    result = await send_interview_reminder(
        user_email=user["email"],
        candidate_name=user["name"],
        company_name=company_name,
        job_role=job_role,
        interview_date=interview_date,
        interview_time=interview_time,
        interview_type=event.get("interview_type", event.get("event_type", "interview")),
        location=event.get("location"),
        meeting_link=event.get("meeting_link"),
        reminder_type="test",
        event_id=event["id"],
        user_id=current_user["id"],
        job_application_id=event.get("job_application_id")
    )
    
    return {"message": "Test reminder sent", "result": result}

@api_router.get("/notifications/logs")
async def get_notification_logs(current_user: dict = Depends(get_current_user)):
    """Get notification logs for the current user"""
    logs = await db.notification_logs.find(
        {"user_id": current_user["id"]}, 
        {"_id": 0}
    ).sort("created_at", -1).to_list(50)
    return logs

@api_router.post("/scheduler/run-check")
async def run_scheduler_check(current_user: dict = Depends(get_current_user)):
    """Manually trigger the reminder scheduler check (for debugging)"""
    logger.info(f"Manual scheduler check triggered by user {current_user['id']}")
    await check_and_send_reminders()
    return {"message": "Scheduler check completed", "check_time": datetime.now(timezone.utc).isoformat()}

@api_router.get("/scheduler/status")
async def get_scheduler_status(current_user: dict = Depends(get_current_user)):
    """Get scheduler status and upcoming events info"""
    now = datetime.now(timezone.utc)
    
    # Get all upcoming events with reminders enabled
    all_events = await db.calendar_events.find({
        "user_id": current_user["id"],
        "reminders_enabled": {"$eq": True},
        "event_type": {"$in": ["interview", "phone_screen", "video_call"]}
    }, {"_id": 0}).to_list(50)
    
    events_info = []
    for event in all_events:
        try:
            event_start = parse_event_datetime(event.get("start_date", ""))
            time_until = event_start - now
            hours_until = time_until.total_seconds() / 3600
            
            events_info.append({
                "id": event.get("id"),
                "title": event.get("title"),
                "start_date": event.get("start_date"),
                "hours_until_event": round(hours_until, 2),
                "reminder_24hr_sent": event.get("reminder_24hr_sent", False),
                "reminder_1hr_sent": event.get("reminder_1hr_sent", False),
                "eligible_24hr": 23 <= hours_until <= 25 and not event.get("reminder_24hr_sent", False),
                "eligible_1hr": 0.833 <= hours_until <= 1.167 and not event.get("reminder_1hr_sent", False)
            })
        except Exception as e:
            events_info.append({
                "id": event.get("id"),
                "title": event.get("title"),
                "error": str(e)
            })
    
    return {
        "server_time_utc": now.isoformat(),
        "scheduler_running": scheduler.running if scheduler else False,
        "scheduler_interval": "5 minutes",
        "upcoming_events": events_info
    }

# ===================== PUBLIC DEBUG ENDPOINT (Development Only) =====================

@api_router.get("/debug/scheduler")
async def debug_scheduler_status():
    """
    PUBLIC debug endpoint for scheduler status - NO AUTHENTICATION REQUIRED.
    Only available when DEBUG_MODE=true in environment.
    Returns: last run time, eligible events count, and execution logs.
    """
    if not DEBUG_MODE:
        raise HTTPException(
            status_code=403, 
            detail="Debug endpoint is disabled in production. Set DEBUG_MODE=true to enable."
        )
    
    now = datetime.now(timezone.utc)
    
    # Get total counts from database
    total_events_with_reminders = await db.calendar_events.count_documents({
        "reminders_enabled": {"$eq": True},
        "event_type": {"$in": ["interview", "phone_screen", "video_call"]}
    })
    
    # Get next scheduler run time
    next_run = None
    if scheduler.running:
        job = scheduler.get_job('reminder_check')
        if job and job.next_run_time:
            next_run = job.next_run_time.isoformat()
    
    # Calculate time since last run
    time_since_last_run = None
    if scheduler_status.get("last_run_time"):
        try:
            last_run = datetime.fromisoformat(scheduler_status["last_run_time"])
            time_since_last_run = f"{int((now - last_run).total_seconds())} seconds ago"
        except:
            pass
    
    return {
        "debug_mode": True,
        "server_time_utc": now.isoformat(),
        "scheduler": {
            "running": scheduler.running if scheduler else False,
            "interval": "5 minutes",
            "next_run_time": next_run,
            "last_run_time": scheduler_status.get("last_run_time"),
            "time_since_last_run": time_since_last_run,
            "last_run_duration_ms": scheduler_status.get("last_run_duration_ms")
        },
        "last_execution": {
            "total_events_checked": scheduler_status.get("total_events_checked", 0),
            "eligible_24hr_count": scheduler_status.get("eligible_24hr_count", 0),
            "eligible_1hr_count": scheduler_status.get("eligible_1hr_count", 0),
            "reminders_sent_24hr": scheduler_status.get("reminders_sent_24hr", 0),
            "reminders_sent_1hr": scheduler_status.get("reminders_sent_1hr", 0)
        },
        "database": {
            "total_events_with_reminders_enabled": total_events_with_reminders
        },
        "errors": scheduler_status.get("last_errors", []),
        "execution_logs": scheduler_status.get("execution_logs", [])
    }

@api_router.post("/debug/scheduler/trigger")
async def debug_trigger_scheduler():
    """
    PUBLIC debug endpoint to manually trigger scheduler - NO AUTHENTICATION REQUIRED.
    Only available when DEBUG_MODE=true in environment.
    """
    if not DEBUG_MODE:
        raise HTTPException(
            status_code=403, 
            detail="Debug endpoint is disabled in production. Set DEBUG_MODE=true to enable."
        )
    
    logger.info("Debug: Manual scheduler trigger initiated")
    await check_and_send_reminders()
    
    return {
        "message": "Scheduler check completed",
        "triggered_at": datetime.now(timezone.utc).isoformat(),
        "status": scheduler_status
    }

# ===================== INTERVIEW PREPARATION ROUTES =====================

@api_router.post("/interview-prep/generate", response_model=InterviewPrepResponse)
async def generate_interview_prep(request: InterviewPrepRequest, current_user: dict = Depends(get_current_user)):
    """Generate comprehensive interview preparation based on job application and resume"""
    
    # Get the job application
    application = await db.applications.find_one(
        {"id": request.application_id, "user_id": current_user["id"]}, 
        {"_id": 0}
    )
    if not application:
        raise HTTPException(status_code=404, detail="Job application not found")
    
    # Get the resume
    resume = await db.resumes.find_one(
        {"id": request.resume_id, "user_id": current_user["id"]}, 
        {"_id": 0}
    )
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")
    
    position = application.get('position', 'the role')
    company = application.get('company', 'the company')
    job_description = application.get('job_description', '')
    resume_content = resume.get('content', '')
    
    # Optionally get the most recent job match analysis for this resume
    match_analysis = None
    match_score = None
    if request.include_match_analysis:
        match = await db.job_matches.find_one(
            {"user_id": current_user["id"], "resume_id": request.resume_id},
            {"_id": 0}
        )
        if match:
            match_analysis = match.get("analysis", {})
            match_score = match_analysis.get("match_score")
    
    # Build context for AI
    weak_areas_context = ""
    if match_analysis:
        missing_skills = match_analysis.get("missing_skills", [])
        weak_areas = match_analysis.get("weak_areas", [])
        if missing_skills or weak_areas:
            weak_areas_context = f"""
Based on previous analysis, the candidate has these gaps:
- Missing Skills: {', '.join(missing_skills) if missing_skills else 'None identified'}
- Weak Areas: {', '.join(weak_areas) if weak_areas else 'None identified'}

Focus interview questions on these areas to help the candidate prepare.
"""
    
    # Try AI generation first
    ai_succeeded = False
    analysis = None
    
    system_msg = """You are an expert interview coach and career advisor. Generate comprehensive, personalized interview preparation materials.

Your response must be practical, specific to the role, and actionable.

Create a structured interview preparation guide with these sections:

1. ROLE OVERVIEW:
   - Brief summary of what the role entails
   - Key responsibilities to understand
   - Skills that will be assessed

2. HR/BEHAVIORAL QUESTIONS (5-7 questions):
   - Questions about teamwork, leadership, conflict resolution, career goals
   - Include STAR method guidance (Situation, Task, Action, Result)
   - Include difficulty level (easy/medium/hard)

3. TECHNICAL QUESTIONS (5-8 questions):
   - Based on job requirements and technical skills needed
   - Range from basic to advanced
   - Include hints and key points for answering

4. SCENARIO-BASED QUESTIONS (3-5 questions):
   - Real-world problem-solving scenarios
   - Role-specific challenges they might face
   - Include approach guidance

5. PROJECT-BASED QUESTIONS (2-3 questions):
   - Questions based on projects mentioned in resume
   - Deep-dive questions about their experience

6. WEAK AREAS TO PREPARE:
   - Topics the candidate should study
   - Specific preparation tips
   - Learning resources

7. PREPARATION TIPS:
   - Interview best practices for this specific role
   - Company-specific advice
   - What to bring/prepare

8. COMPANY RESEARCH POINTS:
   - What to research about the company
   - Industry trends to know

9. QUESTIONS TO ASK THE INTERVIEWER:
   - Smart questions showing interest and preparation

Respond ONLY in this exact JSON format:
{
    "role_overview": {
        "summary": "Brief role summary",
        "key_responsibilities": ["Responsibility 1", "Responsibility 2"],
        "skills_assessed": ["Skill 1", "Skill 2"]
    },
    "hr_behavioral_questions": [
        {
            "question": "Tell me about a time...",
            "category": "hr_behavioral",
            "difficulty": "medium",
            "guidance": ["Use STAR method", "Focus on your specific role", "Quantify results"],
            "sample_points": ["Key point 1", "Key point 2"]
        }
    ],
    "technical_questions": [
        {
            "question": "How would you...",
            "category": "technical",
            "difficulty": "medium",
            "guidance": ["Explain concept first", "Give practical example"],
            "sample_points": ["Technical point 1", "Technical point 2"]
        }
    ],
    "scenario_questions": [
        {
            "question": "Imagine you are...",
            "category": "scenario",
            "difficulty": "hard",
            "guidance": ["Break down the problem", "Consider stakeholders"],
            "sample_points": ["Approach step 1", "Approach step 2"]
        }
    ],
    "project_questions": [
        {
            "question": "Tell me about [project from resume]...",
            "category": "project",
            "difficulty": "medium",
            "guidance": ["Explain your role", "Discuss challenges"],
            "sample_points": ["Your contribution", "Results achieved"]
        }
    ],
    "weak_areas": [
        {
            "topic": "Topic name",
            "reason": "Why this needs preparation",
            "preparation_tips": ["Tip 1", "Tip 2"],
            "resources": ["Resource 1", "Resource 2"]
        }
    ],
    "general_tips": ["Tip 1", "Tip 2"],
    "company_research_points": ["Research point 1", "Research point 2"],
    "questions_to_ask": ["Question 1", "Question 2"]
}"""

    user_msg = f"""Generate comprehensive interview preparation for:

POSITION: {position}
COMPANY: {company}

JOB DESCRIPTION:
{job_description if job_description else 'No specific job description provided - generate general preparation for this role type'}

CANDIDATE'S RESUME:
{resume_content if resume_content else 'No resume content provided'}

{weak_areas_context}

Create personalized, role-specific interview preparation materials that will help this candidate succeed."""

    # Try to get AI response
    response, ai_succeeded = await get_llm_response_safe(system_msg, user_msg)
    
    if ai_succeeded and response:
        try:
            import json
            clean_response = response.strip()
            if clean_response.startswith("```json"):
                clean_response = clean_response[7:]
            if clean_response.startswith("```"):
                clean_response = clean_response[3:]
            if clean_response.endswith("```"):
                clean_response = clean_response[:-3]
            prep_data = json.loads(clean_response.strip())
            
            # Extract role overview if present
            role_overview = prep_data.get("role_overview", {})
            
            # Parse questions, including new project_questions
            all_questions = []
            all_questions.extend([InterviewQuestion(**q) for q in prep_data.get("hr_behavioral_questions", [])])
            all_questions.extend([InterviewQuestion(**q) for q in prep_data.get("technical_questions", [])])
            all_questions.extend([InterviewQuestion(**q) for q in prep_data.get("scenario_questions", [])])
            all_questions.extend([InterviewQuestion(**q) for q in prep_data.get("project_questions", [])])
            
            # Parse and validate the response
            analysis = InterviewPrepAnalysis(
                hr_behavioral_questions=[InterviewQuestion(**q) for q in prep_data.get("hr_behavioral_questions", [])],
                technical_questions=[InterviewQuestion(**q) for q in prep_data.get("technical_questions", [])],
                scenario_questions=[InterviewQuestion(**q) for q in prep_data.get("scenario_questions", [])],
                weak_areas=[WeakArea(**w) for w in prep_data.get("weak_areas", [])],
                general_tips=prep_data.get("general_tips", []),
                company_research_points=prep_data.get("company_research_points", []),
                questions_to_ask=prep_data.get("questions_to_ask", [])
            )
            logger.info(f"AI interview prep generated successfully for {position} at {company}")
        except Exception as e:
            logger.error(f"Failed to parse AI interview prep response: {e}")
            ai_succeeded = False
    
    # If AI failed, use comprehensive fallback
    if not ai_succeeded or not analysis:
        logger.warning(f"Using fallback interview prep for {position} at {company}")
        analysis = generate_fallback_interview_prep(position, company, job_description, resume_content)
    
    # Store in database
    prep_id = str(uuid.uuid4())
    now = datetime.now(timezone.utc).isoformat()
    
    prep_doc = {
        "id": prep_id,
        "user_id": current_user["id"],
        "application_id": request.application_id,
        "resume_id": request.resume_id,
        "job_title": position,
        "company_name": company,
        "analysis": analysis.model_dump(),
        "match_score": match_score,
        "ai_generated": ai_succeeded,
        "created_at": now,
        "updated_at": now
    }
    
    await db.interview_preps.insert_one(prep_doc)
    logger.info(f"Interview prep {'(AI)' if ai_succeeded else '(fallback)'} saved: {prep_id} for user {current_user['id']}")
    
    return InterviewPrepResponse(**{k: v for k, v in prep_doc.items() if k != "_id"})


def generate_fallback_interview_prep(position: str, company: str, job_description: str, resume_content: str) -> InterviewPrepAnalysis:
    """
    Generate a comprehensive fallback interview preparation guide.
    This is used when AI service is unavailable.
    """
    
    # Extract some keywords from job description for personalization
    tech_keywords = []
    if job_description:
        common_tech = ['python', 'javascript', 'react', 'node', 'sql', 'aws', 'docker', 'kubernetes', 
                       'java', 'c++', 'golang', 'rust', 'typescript', 'angular', 'vue', 'mongodb',
                       'postgresql', 'redis', 'kafka', 'spark', 'machine learning', 'ai', 'data']
        job_desc_lower = job_description.lower()
        tech_keywords = [tech for tech in common_tech if tech in job_desc_lower][:5]
    
    # Build comprehensive fallback content
    hr_behavioral_questions = [
        InterviewQuestion(
            question=f"Tell me about yourself and why you're interested in the {position} role at {company}.",
            category="hr_behavioral",
            difficulty="easy",
            guidance=["Keep it professional and concise (2-3 minutes)", "Connect your background to the role", "Show enthusiasm for the company"],
            sample_points=["Your professional background", "Relevant experience and skills", "Why this role excites you"]
        ),
        InterviewQuestion(
            question="Describe a challenging project you worked on. What was your role and how did you handle obstacles?",
            category="hr_behavioral",
            difficulty="medium",
            guidance=["Use the STAR method (Situation, Task, Action, Result)", "Focus on your specific contributions", "Quantify results where possible"],
            sample_points=["The challenge you faced", "Actions you took", "Measurable outcomes achieved"]
        ),
        InterviewQuestion(
            question="Tell me about a time when you had a conflict with a colleague. How did you resolve it?",
            category="hr_behavioral",
            difficulty="medium",
            guidance=["Show emotional intelligence", "Focus on resolution, not the conflict", "Demonstrate professionalism"],
            sample_points=["How you approached the conversation", "Steps to find common ground", "Positive outcome achieved"]
        ),
        InterviewQuestion(
            question="Where do you see yourself in 5 years? How does this role fit into your career goals?",
            category="hr_behavioral",
            difficulty="easy",
            guidance=["Show ambition but be realistic", "Align goals with company growth", "Demonstrate commitment"],
            sample_points=["Skills you want to develop", "Leadership aspirations", "Industry expertise goals"]
        ),
        InterviewQuestion(
            question="Describe a time when you had to learn something new quickly. How did you approach it?",
            category="hr_behavioral",
            difficulty="medium",
            guidance=["Show adaptability and learning agility", "Describe your learning process", "Highlight the outcome"],
            sample_points=["Your learning strategy", "Resources you used", "How you applied the knowledge"]
        ),
        InterviewQuestion(
            question="Tell me about a time you received constructive criticism. How did you handle it?",
            category="hr_behavioral",
            difficulty="medium",
            guidance=["Show growth mindset", "Demonstrate self-awareness", "Focus on improvement"],
            sample_points=["How you received the feedback", "Actions you took to improve", "Results of your changes"]
        )
    ]
    
    # Technical questions based on role or generic
    technical_questions = [
        InterviewQuestion(
            question=f"Walk me through your technical background and how it prepares you for this {position} role.",
            category="technical",
            difficulty="easy",
            guidance=["Start with your strongest skills", "Connect to job requirements", "Give specific examples"],
            sample_points=["Core technical competencies", "Relevant projects", "Continuous learning efforts"]
        ),
        InterviewQuestion(
            question="Describe your approach to debugging a complex issue in production.",
            category="technical",
            difficulty="medium",
            guidance=["Show systematic thinking", "Mention tools you use", "Discuss prevention strategies"],
            sample_points=["Initial diagnosis steps", "Root cause analysis", "Long-term fixes implemented"]
        ),
        InterviewQuestion(
            question="How do you ensure code quality in your work? What practices do you follow?",
            category="technical",
            difficulty="medium",
            guidance=["Mention testing strategies", "Discuss code reviews", "Talk about documentation"],
            sample_points=["Testing approaches (unit, integration)", "Code review process", "Best practices you follow"]
        ),
        InterviewQuestion(
            question="Explain a technical concept you know well to me as if I were a non-technical stakeholder.",
            category="technical",
            difficulty="medium",
            guidance=["Use analogies and simple language", "Focus on business value", "Avoid jargon"],
            sample_points=["Clear explanation", "Real-world comparison", "Why it matters"]
        ),
        InterviewQuestion(
            question="How do you stay current with new technologies and industry trends?",
            category="technical",
            difficulty="easy",
            guidance=["Show passion for learning", "Mention specific resources", "Discuss practical application"],
            sample_points=["Learning resources you use", "Communities you're part of", "Recent things you've learned"]
        )
    ]
    
    # Add tech-specific questions if keywords found
    if tech_keywords:
        tech_list = ', '.join(tech_keywords[:3])
        technical_questions.append(
            InterviewQuestion(
                question=f"I see this role involves {tech_list}. Describe your experience with these technologies.",
                category="technical",
                difficulty="medium",
                guidance=["Be specific about your experience level", "Give project examples", "Mention challenges overcome"],
                sample_points=["Projects using these technologies", "Proficiency level", "Best practices you follow"]
            )
        )
    
    scenario_questions = [
        InterviewQuestion(
            question=f"Imagine you join {company} and are given a project with unclear requirements. How would you proceed?",
            category="scenario",
            difficulty="medium",
            guidance=["Show initiative and communication skills", "Discuss stakeholder management", "Mention risk mitigation"],
            sample_points=["Clarification questions you'd ask", "How you'd document requirements", "Iterative approach"]
        ),
        InterviewQuestion(
            question="You discover a critical bug just before a major release. Walk me through your decision-making process.",
            category="scenario",
            difficulty="hard",
            guidance=["Show prioritization skills", "Discuss communication approach", "Balance urgency with quality"],
            sample_points=["Immediate assessment steps", "Stakeholder communication", "Risk evaluation"]
        ),
        InterviewQuestion(
            question="A team member is consistently missing deadlines, affecting your work. How do you handle this?",
            category="scenario",
            difficulty="medium",
            guidance=["Show empathy and professionalism", "Focus on solutions", "Know when to escalate"],
            sample_points=["Private conversation approach", "Understanding root causes", "Collaborative solutions"]
        ),
        InterviewQuestion(
            question="You're asked to implement a feature you believe is technically flawed. What do you do?",
            category="scenario",
            difficulty="hard",
            guidance=["Balance technical opinion with business needs", "Communicate concerns professionally", "Propose alternatives"],
            sample_points=["How you'd voice concerns", "Alternative solutions", "Documentation of risks"]
        )
    ]
    
    weak_areas = [
        WeakArea(
            topic="Company Knowledge",
            reason=f"Understanding {company}'s products, culture, and recent developments will show genuine interest",
            preparation_tips=[
                f"Research {company}'s website, especially About Us and Products pages",
                "Look up recent news articles and press releases",
                "Check LinkedIn for company updates and employee posts",
                "Understand their competitors and market position"
            ],
            resources=["Company website", "LinkedIn company page", "Glassdoor reviews", "Recent news articles"]
        ),
        WeakArea(
            topic="Role-Specific Requirements",
            reason=f"Deep understanding of {position} responsibilities will help you give relevant answers",
            preparation_tips=[
                "Review the job description multiple times",
                "Research similar roles at other companies",
                "Identify key skills and prepare examples for each",
                "Understand common challenges in this role"
            ],
            resources=["Job posting", "Industry blogs", "Professional forums", "LinkedIn job descriptions"]
        ),
        WeakArea(
            topic="STAR Method Stories",
            reason="Behavioral questions require structured, specific stories from your experience",
            preparation_tips=[
                "Prepare 5-7 detailed stories using STAR format",
                "Include stories about leadership, conflict, failure, and success",
                "Practice telling each story in 2-3 minutes",
                "Have metrics and specific outcomes ready"
            ],
            resources=["Your resume and past projects", "Performance reviews", "Colleague feedback"]
        )
    ]
    
    general_tips = [
        f"Research {company} thoroughly - know their products, mission, and recent news",
        "Prepare specific examples from your experience using the STAR method",
        "Practice answering questions out loud, not just in your head",
        "Prepare thoughtful questions to ask the interviewer",
        "Dress professionally and arrive 10-15 minutes early",
        "Bring copies of your resume and a notepad",
        "Follow up with a thank-you email within 24 hours",
        "Get a good night's sleep before the interview",
        "Review the job description one more time before the interview",
        "Be ready to discuss salary expectations if asked"
    ]
    
    company_research_points = [
        f"{company}'s mission statement and core values",
        "Recent company news, product launches, or acquisitions",
        "Company culture and work environment",
        "Key competitors and market position",
        "Growth trajectory and future plans",
        "Leadership team and organizational structure",
        "Employee reviews on Glassdoor and LinkedIn"
    ]
    
    questions_to_ask = [
        f"What does success look like for a {position} in the first 90 days?",
        "How would you describe the team culture?",
        "What are the biggest challenges facing the team right now?",
        "How does this role contribute to the company's overall goals?",
        "What opportunities for growth and development does the company offer?",
        "Can you describe a typical day or week in this role?",
        "What do you enjoy most about working at " + company + "?",
        "What are the next steps in the interview process?"
    ]
    
    return InterviewPrepAnalysis(
        hr_behavioral_questions=hr_behavioral_questions,
        technical_questions=technical_questions,
        scenario_questions=scenario_questions,
        weak_areas=weak_areas,
        general_tips=general_tips,
        company_research_points=company_research_points,
        questions_to_ask=questions_to_ask
    )

@api_router.get("/interview-prep", response_model=List[InterviewPrepResponse])
async def get_interview_preps(application_id: Optional[str] = None, current_user: dict = Depends(get_current_user)):
    """Get all interview preparations for the user, optionally filtered by application"""
    query = {"user_id": current_user["id"]}
    if application_id:
        query["application_id"] = application_id
    
    preps = await db.interview_preps.find(query, {"_id": 0}).sort("created_at", -1).to_list(100)
    return [InterviewPrepResponse(**p) for p in preps]

@api_router.get("/interview-prep/{prep_id}", response_model=InterviewPrepResponse)
async def get_interview_prep(prep_id: str, current_user: dict = Depends(get_current_user)):
    """Get a specific interview preparation"""
    prep = await db.interview_preps.find_one({"id": prep_id, "user_id": current_user["id"]}, {"_id": 0})
    if not prep:
        raise HTTPException(status_code=404, detail="Interview preparation not found")
    return InterviewPrepResponse(**prep)

@api_router.delete("/interview-prep/{prep_id}")
async def delete_interview_prep(prep_id: str, current_user: dict = Depends(get_current_user)):
    """Delete an interview preparation"""
    result = await db.interview_preps.delete_one({"id": prep_id, "user_id": current_user["id"]})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Interview preparation not found")
    return {"message": "Interview preparation deleted"}

@api_router.post("/interview-prep/{prep_id}/regenerate", response_model=InterviewPrepResponse)
async def regenerate_interview_prep(prep_id: str, current_user: dict = Depends(get_current_user)):
    """Regenerate interview preparation with fresh questions"""
    existing = await db.interview_preps.find_one({"id": prep_id, "user_id": current_user["id"]}, {"_id": 0})
    if not existing:
        raise HTTPException(status_code=404, detail="Interview preparation not found")
    
    # Delete old and create new
    await db.interview_preps.delete_one({"id": prep_id})
    
    request = InterviewPrepRequest(
        application_id=existing["application_id"],
        resume_id=existing["resume_id"],
        include_match_analysis=True
    )
    
    return await generate_interview_prep(request, current_user)

# ===================== ANALYTICS ROUTES =====================

@api_router.get("/analytics")
async def get_analytics(current_user: dict = Depends(get_current_user)):
    user_id = current_user["id"]
    
    # Get application stats
    pipeline = [
        {"$match": {"user_id": user_id}},
        {"$group": {"_id": "$status", "count": {"$sum": 1}}}
    ]
    status_counts = await db.applications.aggregate(pipeline).to_list(20)
    
    total_apps = sum(s["count"] for s in status_counts)
    status_breakdown = {s["_id"]: s["count"] for s in status_counts}
    
    # Get recent applications
    recent_apps = await db.applications.find({"user_id": user_id}, {"_id": 0}).sort("created_at", -1).limit(5).to_list(5)
    
    # Get resume count
    resume_count = await db.resumes.count_documents({"user_id": user_id})
    
    # Get cover letter count (from both legacy and v2 collections)
    cover_letter_count_legacy = await db.cover_letters.count_documents({"user_id": user_id})
    cover_letter_count_v2 = await db.cover_letters_v2.count_documents({"user_id": user_id})
    cover_letter_count = cover_letter_count_legacy + cover_letter_count_v2
    
    # Get upcoming events
    now = datetime.now(timezone.utc).isoformat()
    upcoming_events = await db.calendar_events.find(
        {"user_id": user_id, "start_date": {"$gte": now}}, 
        {"_id": 0}
    ).sort("start_date", 1).limit(5).to_list(5)
    
    # Calculate response rate
    responded_statuses = ["interviewing", "offer", "rejected"]
    responded = sum(status_breakdown.get(s, 0) for s in responded_statuses)
    response_rate = round((responded / total_apps * 100) if total_apps > 0 else 0, 1)
    
    return {
        "total_applications": total_apps,
        "status_breakdown": status_breakdown,
        "response_rate": response_rate,
        "resume_count": resume_count,
        "cover_letter_count": cover_letter_count,
        "recent_applications": recent_apps,
        "upcoming_events": upcoming_events
    }

# ===================== HEALTH CHECK =====================

@api_router.get("/")
async def root():
    return {"message": "HireFlow AI API is running", "version": "1.0.0"}

@api_router.get("/health")
async def health_check():
    return {"status": "healthy", "timestamp": datetime.now(timezone.utc).isoformat()}

# Include router and add middleware
app.include_router(api_router)

app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=os.environ.get('CORS_ORIGINS', '*').split(','),
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.on_event("startup")
async def startup_event():
    """Start background scheduler for email reminders"""
    if RESEND_API_KEY:
        # Schedule reminder check every 5 minutes
        scheduler.add_job(
            check_and_send_reminders,
            'interval',
            minutes=5,
            id='reminder_check',
            replace_existing=True,
            next_run_time=datetime.now(timezone.utc) + timedelta(seconds=30)  # First run in 30 seconds
        )
        scheduler.start()
        logger.info("=" * 50)
        logger.info("Email reminder scheduler started")
        logger.info(f"  - Check interval: every 5 minutes")
        logger.info(f"  - First check in: 30 seconds")
        logger.info(f"  - Server time (UTC): {datetime.now(timezone.utc).isoformat()}")
        logger.info("=" * 50)
    else:
        logger.warning("RESEND_API_KEY not configured - email reminders disabled")

@app.on_event("shutdown")
async def shutdown_db_client():
    if scheduler.running:
        scheduler.shutdown()
    client.close()
